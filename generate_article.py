"""Generate LinkedIn article as .docx for the Ear Landmarker project."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BLAZEEAR_URL = (
    "https://www.linkedin.com/pulse/blazeear-extending-trainable-blazeface-"
    "from-faces-ears-shameem-hameed-dalic/"
)
BLAZEFACE_URL = (
    "https://www.linkedin.com/pulse/trainable-blazeface-high-performance-"
    "extendable-face-shameem-hameed-jvxyc/"
)
HEADSIZE_URL = (
    "https://www.linkedin.com/pulse/measuring-human-face-privately-browser-"
    "shameem-hameed-51okc/"
)
DEMO_URL = "https://shameem4.github.io/Ear_Landmarker/"
GITHUB_URL = "https://github.com/shameem4/Ear_Landmarker"


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph (python-docx lacks built-in support)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")  # 11pt
    rPr.append(sz)
    run_elem.append(rPr)
    run_elem.text = text
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_para_with_links(doc, segments):
    """Add a paragraph with mixed plain text and hyperlinks.

    segments: list of (text, url_or_None) tuples.
    """
    p = doc.add_paragraph()
    for text, url in segments:
        if url:
            add_hyperlink(p, text, url)
        else:
            run = p.add_run(text)
            run.font.size = Pt(11)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    """Render table as ASCII text in monospace font for clean copy-paste."""
    all_rows = [headers] + [[str(v) for v in r] for r in rows]
    col_widths = [
        max(len(row[c]) for row in all_rows) for c in range(len(headers))
    ]

    def fmt_row(row):
        cells = [row[c].ljust(col_widths[c]) for c in range(len(headers))]
        return "  " + "  |  ".join(cells)

    def separator():
        return "  " + "--+--".join("-" * w for w in col_widths)

    lines = [fmt_row(all_rows[0]), separator()]
    for row in all_rows[1:]:
        lines.append(fmt_row(row))

    p = doc.add_paragraph()
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
        if i < len(lines) - 1:
            p.add_run("\n").font.size = Pt(9)

    doc.add_paragraph()  # spacing


def build_article():
    doc = Document()

    # Adjust default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ===== Title =====
    title = doc.add_heading(
        "Real-Time 55-Point Ear Landmark Detection: "
        "A BlazeBlock Pipeline from Data Curation to Browser Deployment",
        level=0,
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== Demo + Repo links =====
    add_para_with_links(doc, [
        ("Try the ", None),
        ("live browser demo", DEMO_URL),
        (" | ", None),
        ("GitHub repository", GITHUB_URL),
    ])

    # ===== The Short Version =====
    add_heading(doc, "The Short Version")
    add_para(
        doc,
        "I built a real-time ear landmark detection system that locates 55 anatomical "
        "points on human ears from a single RGB image. The pipeline runs two stages: "
        "a BlazeEar detector finds ears in the frame, then an EarLandmarker network "
        "regresses 55 landmarks on each cropped ear region."
    )
    add_para(
        doc,
        "4 datasets. 5,870 samples. 3 annotation formats. 312K parameters. "
        "Validation NME of 0.0307 (approximately 5.9 pixels at 192x192 resolution). "
        "579 FPS on GPU. 226 FPS on CPU. Runs in the browser via ONNX Runtime Web."
    )
    add_para_with_links(doc, [
        ("This is the third project in a series. The first, ", None),
        ("Trainable BlazeFace", BLAZEFACE_URL),
        (", rebuilt MediaPipe's face detector as a fully trainable PyTorch model "
         "with browser deployment. The second, ", None),
        ("BlazeEar", BLAZEEAR_URL),
        (", extended that architecture from faces to ears. This project adds "
         "landmark regression on top of the BlazeEar detector, completing the "
         "pipeline from raw image to 55 anatomical coordinates.", None),
    ])
    add_para(
        doc,
        "Here is what the system does, how the data was assembled, what architectural "
        "choices were made and why, and where this is heading."
    )

    # ===== Why This Matters =====
    add_heading(doc, "Why This Matters")
    add_para(
        doc,
        "Ear biometrics are an underexplored modality in computer vision. Ears are "
        "visible in profile and oblique views where faces are partially or fully "
        "occluded. They do not change expression. Their anatomy is stable across "
        "decades. Unlike fingerprints, they can be captured at a distance without "
        "subject cooperation. These properties make ears valuable for identification, "
        "authentication, and human-computer interaction."
    )
    add_para(
        doc,
        "Ear landmark detection is the prerequisite for all downstream ear analysis: "
        "biometric identification, ear shape classification, hearing aid fitting, "
        "personalized spatial audio (HRTF estimation), and clinical morphology "
        "assessment. Without reliable landmarks, none of these applications can "
        "operate on unconstrained images."
    )
    add_para(
        doc,
        "The challenge is data scarcity. Ear landmark datasets are small, "
        "fragmented across incompatible annotation formats, and rarely quality-controlled. "
        "No single public dataset exceeds 6,500 samples. Most models in the literature "
        "train and evaluate on a single source, making cross-dataset generalization unknown."
    )
    add_para(
        doc,
        "The goal of this project: unify all available 55-point ear landmark datasets "
        "into one clean corpus, train a lightweight model suitable for real-time inference "
        "on consumer hardware, and deploy it end-to-end in the browser with no server-side "
        "computation."
    )

    # ===== The Data =====
    add_heading(doc, "The Data")
    add_para(
        doc,
        "I assembled four publicly available ear landmark datasets spanning three "
        "annotation formats:"
    )
    add_table(doc,
        ["Source", "Samples", "Format", "Resolution", "Notes"],
        [
            ["Collection B", "3,153", ".pts (55-point)", "256x256", "Pre-cropped ear images"],
            ["AudioEar2D", "2,000", "LabelMe JSON", "299x299", "Synthetic ears from FFHQ, StyleGAN-generated"],
            ["Collection A", "605", ".pts (55-point)", "Variable", "Full-face images, auto-cropped to ear ROI"],
            ["AudioEar3D", "112", "LabelMe JSON", "187x186", "3D-scanned ear renders"],
        ],
    )
    add_para(
        doc,
        "Total: 5,870 unique samples after deduplication by image hash and removal "
        "of flipped variants (flipping is handled in augmentation instead). All landmarks "
        "normalized to [0, 1] coordinate space."
    )

    add_para(doc, "Data provenance and quality control", bold=True)
    add_para(
        doc,
        "A fifth dataset (COCO-format keypoints, 629 images) was evaluated and dropped. "
        "Analysis of inter-point step distances revealed that its landmark ordering does not "
        "match the linestrip convention (helix 0-19, antihelix 20-34, concha 35-49, "
        "tragus 50-54) used by the other four sources. Mean step distances were 2-4x higher "
        "than expected, confirming an incompatible annotation scheme. No reliable mapping "
        "was found. Including it degraded training."
    )
    add_para(
        doc,
        "Automated quality validation runs seven checks on the unified dataset: image "
        "integrity (can every image be decoded), landmark bounds (points outside [0, 1] or "
        "clipped to edges), bounding box geometry (degenerate shapes), linestrip smoothness "
        "(large jumps between consecutive points suggesting misordering), per-point statistical "
        "outliers (>3 sigma from the per-landmark mean), cross-source mean shape consistency "
        "(Procrustes-like distance between source-level average shapes), and resolution audit "
        "(extreme aspect ratios, images below 32 pixels)."
    )
    add_para(
        doc,
        "Stratified train/validation split (85/15) ensures each source is represented "
        "proportionally in both sets. 4,987 training samples, 877 validation samples. "
        "No test set is held out because the total corpus is small enough that maximizing "
        "training data is preferred; evaluation is on the validation split."
    )

    # ===== Architecture =====
    add_heading(doc, "Architecture")
    add_para_with_links(doc, [
        ("The system uses a two-stage pipeline at inference time. A ", None),
        ("BlazeEar detector", BLAZEEAR_URL),
        (" (described in a previous article) locates ears in the input frame. "
         "Each detected bounding box is expanded by 30% for spatial context, cropped, "
         "and resized to 192x192 pixels. The EarLandmarker network then regresses "
         "55 (x, y) coordinates from this crop.", None),
    ])

    add_para(doc, "EarLandmarker (312,958 parameters)", bold=True)
    add_para(
        doc,
        "The architecture follows the MediaPipe FaceMesh pattern: a BlazeBlock backbone "
        "with progressive channel expansion, followed by global average pooling and a "
        "fully connected regression head with sigmoid activation. The output is 110 values "
        "(55 landmarks x 2 coordinates), each in [0, 1], representing normalized positions "
        "within the 192x192 crop."
    )
    add_table(doc,
        ["Stage", "Channels", "Spatial", "Blocks", "Parameters"],
        [
            ["conv0", "3 -> 24", "192 -> 96", "1 (Conv 5x5 s=2)", "1,848"],
            ["stage0", "24 -> 24", "96 (no downsample)", "2 BlazeBlocks", "2,544"],
            ["stage1", "24 -> 48", "96 -> 48", "4 BlazeBlocks", "14,232"],
            ["stage2", "48 -> 96", "48 -> 24", "4 BlazeBlocks", "46,896"],
            ["stage3", "96 -> 128", "24 -> 12", "4 BlazeBlocks", "87,968"],
            ["stage4", "128 -> 192", "12 -> 6", "3 BlazeBlocks", "138,240"],
            ["head", "192 -> 110", "GAP + FC + Sigmoid", "-", "21,230"],
        ],
    )

    add_para(doc, "BlazeBlock design", bold=True)
    add_para_with_links(doc, [
        ("Each BlazeBlock is a depthwise separable convolution with a skip connection: "
         "DepthwiseConv (5x5) -> BatchNorm -> PointwiseConv (1x1) -> BatchNorm -> Add -> ReLU. "
         "Stride-2 blocks use max pooling on the skip path. When input and output channels "
         "differ, the skip connection includes a learned 1x1 projection. This is the same "
         "block used in ", None),
        ("Trainable BlazeFace", BLAZEFACE_URL),
        (" and ", None),
        ("BlazeEar", BLAZEEAR_URL),
        (", enabling weight transfer from the detector backbone.", None),
    ])
    add_para(
        doc,
        "TFLite-compatible asymmetric padding is used in the initial convolution and all "
        "stride-2 BlazeBlocks. This matches the MediaPipe convention and ensures the model "
        "can be exported to TFLite or ONNX without padding mismatches."
    )

    add_para(doc, "Backbone transfer", bold=True)
    add_para_with_links(doc, [
        ("The initial conv0 layer and early BlazeBlocks share the same channel dimensions "
         "as the ", None),
        ("BlazeEar", BLAZEEAR_URL),
        (" backbone. A transfer utility loads compatible weights from a trained "
         "BlazeEar checkpoint, providing a warm start. The channel dimensions diverge after "
         "stage0 (BlazeEar goes to 28 channels; the landmarker goes to 48), so transfer is "
         "limited to the first few layers.", None),
    ])

    # ===== Landmark Layout =====
    add_heading(doc, "Landmark Layout")
    add_para(
        doc,
        "The 55 landmarks are organized as four linestrips tracing the major anatomical "
        "structures of the ear:"
    )
    add_table(doc,
        ["Group", "Indices", "Points", "Anatomical Structure"],
        [
            ["Helix", "0-19", "20", "Outer rim of the ear, from superior to inferior"],
            ["Antihelix", "20-34", "15", "Inner ridge, including superior and inferior crus"],
            ["Concha", "35-49", "15", "Bowl-shaped cavity surrounding the ear canal"],
            ["Tragus", "50-54", "5", "Small cartilage flap anterior to the ear canal"],
        ],
    )
    add_para(
        doc,
        "This ordering is consistent across Collection A, Collection B, AudioEar2D, "
        "and AudioEar3D. The linestrip convention means consecutive points within each "
        "group are spatially adjacent, enabling smoothness checks during validation and "
        "connected-line visualization at inference."
    )

    # ===== Training =====
    add_heading(doc, "Training")
    add_table(doc,
        ["Setting", "Value"],
        [
            ["Loss function", "Wing loss (w=0.04, epsilon=0.01)"],
            ["Optimizer", "AdamW (lr=1e-3, weight decay=1e-4)"],
            ["Learning rate schedule", "Cosine annealing to 1e-5"],
            ["Precision", "16-mixed AMP"],
            ["Batch size", "64"],
            ["Input normalization", "[-1, 1] (pixel / 255, subtract 0.5, divide 0.5)"],
            ["Augmentation", "Horizontal flip (50%), translation (5%), rotation (+/-15 deg), "
                            "color jitter (B=0.3, C=0.3, S=0.2, H=0.05), bbox jitter (10%)"],
            ["Early stopping", "Patience 50 epochs on val NME"],
            ["Hardware", "RTX 5060 Ti 16GB, Ryzen 9 9950X3D, 64GB RAM"],
        ],
    )

    add_para(doc, "Wing loss", bold=True)
    add_para(
        doc,
        "Wing loss (Feng et al., CVPR 2018) uses a logarithmic term for small errors and "
        "a linear term for large errors. The transition at w=0.04 means errors below 4% of "
        "the normalized coordinate range receive amplified gradients, prioritizing fine "
        "alignment over gross localization. This is standard for facial and ear landmark "
        "regression -- it outperforms L1 and L2 losses because small positional errors are "
        "perceptually important for landmark precision."
    )

    add_para(doc, "Augmentation rationale", bold=True)
    add_para(
        doc,
        "Bbox jitter (random scale and translation of the crop) simulates variance in the "
        "upstream detector's bounding box predictions. At inference, the detector crop will "
        "not perfectly center the ear. Training with jittered crops makes the landmarker "
        "robust to this. Horizontal flip is applied with mirrored x-coordinates. Rotation "
        "and translation augmentations use affine transforms applied jointly to the image "
        "and landmark coordinates. Color jitter is split: brightness, contrast, and saturation "
        "are fast tensor operations; hue jitter requires an expensive HSV conversion and is "
        "applied only 30% of the time."
    )

    add_para(doc, "Training trajectory", bold=True)
    add_para(
        doc,
        "Baseline training (100 epochs) reached val NME 0.0351 (approximately 6.7 pixels). "
        "Extended training (500 epoch budget, early stopped at epoch 395) reached val NME "
        "0.0307 (approximately 5.9 pixels), a 12.6% improvement. The train-val NME gap at "
        "convergence was 0.0014, indicating no overfitting. The best checkpoint was saved at "
        "epoch 345."
    )

    # ===== Results =====
    add_heading(doc, "Results")
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Val NME (normalized mean error)", "0.0307"],
            ["Approximate pixel error at 192x192", "5.9 px"],
            ["Train-val NME gap", "0.0014"],
            ["GPU inference (RTX 5060 Ti)", "579 FPS"],
            ["CPU inference (Ryzen 9 9950X3D)", "226 FPS"],
            ["Model parameters", "312,958"],
            ["Model size (ONNX, simplified)", "1.2 MB"],
            ["Input resolution", "192 x 192"],
        ],
    )
    add_para(
        doc,
        "NME (Normalized Mean Error) is the mean Euclidean distance between predicted "
        "and ground truth landmarks, normalized by the [0, 1] coordinate range. An NME of "
        "0.0307 means the average landmark is 3.07% of the image width/height from its "
        "ground truth position."
    )

    # ===== Deployment =====
    add_heading(doc, "Deployment")
    add_para(
        doc,
        "The trained model is exported to ONNX (opset 14) with graph simplification "
        "via onnx-simplifier. The resulting 1.2 MB file runs in the browser using "
        "ONNX Runtime Web (WASM backend). A JavaScript inference module implements "
        "the full pipeline: BlazeEar detection (516 KB ONNX), NMS, ROI crop with "
        "30% expansion, EarLandmarker regression, and coordinate mapping back to "
        "frame pixels."
    )
    add_para_with_links(doc, [
        ("The browser demo provides webcam and image upload modes with real-time "
         "visualization of bounding boxes and color-coded linestrip landmarks. "
         "This follows the same deployment pattern used for the ", None),
        ("BlazeEar", BLAZEEAR_URL),
        (" and ", None),
        ("BlazeFace", BLAZEFACE_URL),
        (" browser demos. No server-side computation is required. The full demo is "
         "a static site deployable via GitHub Pages: one HTML file, one JavaScript "
         "module, and two ONNX model files. ", None),
        ("Try it here", DEMO_URL),
        (". Full source code is on ", None),
        ("GitHub", GITHUB_URL),
        (".", None),
    ])

    # ===== Design Decisions =====
    add_heading(doc, "Design Decisions and Trade-offs")

    add_para(doc, "Why BlazeBlocks instead of a standard ResNet or MobileNet backbone", bold=True)
    add_para(
        doc,
        "BlazeBlocks are designed for real-time inference on mobile and edge devices. "
        "The depthwise separable structure keeps the parameter count low (312K vs "
        "millions for ResNet-based approaches) while maintaining sufficient representational "
        "capacity for 55-point regression. The architecture matches the upstream BlazeEar "
        "detector, enabling weight transfer for the early layers. MediaPipe's FaceMesh "
        "uses the same block design to regress 468 facial landmarks in real time -- "
        "55 ear landmarks is a strictly easier task."
    )

    add_para(doc, "Why direct coordinate regression instead of heatmap regression", bold=True)
    add_para(
        doc,
        "Heatmap-based methods (predicting a 2D Gaussian at each landmark location) are "
        "standard in body pose estimation but add substantial computational cost: the "
        "decoder must produce K spatial maps at the output resolution. For 55 landmarks "
        "at 192x192, this would be 55 x 48 x 48 = 126,720 output values (assuming 4x "
        "downsampled heatmaps) versus 110 values for direct regression. Direct regression "
        "also avoids the argmax discretization that limits heatmap precision to the spatial "
        "resolution of the output grid. The trade-off is that direct regression requires "
        "careful loss design (hence Wing loss) to handle the non-uniform error landscape."
    )

    add_para(doc, "Why sigmoid output instead of unbounded regression", bold=True)
    add_para(
        doc,
        "Landmark coordinates are inherently bounded to the image. Sigmoid activation "
        "on the output layer constrains predictions to [0, 1], eliminating the possibility "
        "of predicting landmarks outside the crop. This is a hard architectural constraint "
        "rather than a soft regularization, and it simplifies the loss landscape."
    )

    add_para(doc, "Why unified preprocessing instead of per-source training", bold=True)
    add_para(
        doc,
        "Training separate models per source would avoid cross-source annotation "
        "inconsistencies but would limit each model to a few hundred to a few thousand "
        "samples. The unified corpus provides 5,870 diverse samples. Cross-source "
        "validation (mean shape L2 distance between sources) confirmed that the four "
        "retained datasets share a consistent annotation convention. The dropped COCO "
        "dataset was the only source with incompatible ordering."
    )

    # ===== Limitations =====
    add_heading(doc, "Limitations")
    add_para(
        doc,
        "The dataset is small by deep learning standards. 5,870 samples is adequate for "
        "a 312K parameter model but leaves limited room for error. The synthetic component "
        "(AudioEar2D, 34% of data) may not fully represent real-world ear appearance "
        "diversity. Ethnic, age, and sex distribution across the combined dataset is "
        "not documented in the source publications, making demographic bias assessment "
        "impossible."
    )
    add_para(
        doc,
        "The 55-point annotation scheme is standardized across the retained sources but "
        "is not an ISO or AASM-equivalent standard. Different research groups may define "
        "ear landmarks differently. The linestrip convention (helix / antihelix / concha / "
        "tragus) was verified empirically through smoothness analysis, not from a formal "
        "specification."
    )
    add_para(
        doc,
        "The pipeline depends on the upstream BlazeEar detector. If the detector misses "
        "an ear or produces a poor crop, the landmarker cannot recover. End-to-end "
        "training (joint detection and landmark regression) was not explored."
    )
    add_para(
        doc,
        "No per-landmark error breakdown is currently computed. Some landmarks (e.g., "
        "tragus points near the ear canal) may have systematically higher error than "
        "others. Reporting per-group NME (helix, antihelix, concha, tragus) would "
        "provide finer diagnostic information."
    )

    # ===== Future Directions =====
    add_heading(doc, "Future Directions")

    add_para(doc, "Personalized spatial audio (HRTF estimation)", bold=True)
    add_para(
        doc,
        "Head-related transfer functions describe how sound is filtered by the "
        "geometry of the ear before reaching the eardrum. They are unique to each "
        "individual and critical for realistic 3D audio rendering. Current HRTF "
        "personalization requires either expensive acoustic measurements in an "
        "anechoic chamber or 3D ear scanning. Ear landmarks from a single photo "
        "could parameterize a predictive HRTF model, enabling personalized spatial "
        "audio from a smartphone camera. The 55-point representation captures the "
        "major structures (helix curvature, concha depth, tragus shape) that "
        "determine high-frequency HRTF notches."
    )

    add_para(doc, "Ear biometric identification", bold=True)
    add_para(
        doc,
        "Landmark-based ear recognition is complementary to appearance-based methods. "
        "A geometric descriptor derived from 55 landmarks (inter-point distances, "
        "angles, curvature) is invariant to lighting and partially invariant to "
        "viewpoint. Fusing geometric and appearance features could improve robustness "
        "for profile-view identification in surveillance or access control."
    )

    add_para(doc, "Hearing aid and earbud fitting", bold=True)
    add_para(
        doc,
        "Custom-fit hearing aids and earbuds currently require physical ear "
        "impressions or 3D scanning. Landmark-based ear shape parameterization from "
        "a photo could enable remote fitting workflows, particularly relevant for "
        "telehealth and direct-to-consumer hearing devices."
    )

    add_para(doc, "Ear sizing from a single photo", bold=True)
    add_para_with_links(doc, [
        ("In a previous project, ", None),
        ("Measuring the Human Face Privately in the Browser", HEADSIZE_URL),
        (", I demonstrated that metric head dimensions (circumference, width, "
         "length) can be estimated from a single webcam image using facial landmarks "
         "and a known reference object, with all computation running client-side. "
         "The same principle extends directly to ears.", None),
    ])
    add_para(
        doc,
        "The 55 ear landmarks provide the geometric basis for estimating physical ear "
        "dimensions: total ear height (helix superior to lobule), ear width (tragus to "
        "helix), concha depth and diameter, and inter-landmark distances that parameterize "
        "the full ear shape. Given a scale reference -- a coin held near the ear, a known "
        "interpupillary distance from the face detector, or a depth sensor -- these "
        "normalized landmark coordinates convert to millimeters."
    )
    add_para(
        doc,
        "Applications are immediate. Hearing aid manufacturers specify size by concha "
        "diameter and canal depth. Earbud tip sizing (S/M/L) maps to concha and ear canal "
        "entry dimensions. Over-ear headphone fit depends on total ear height and width. "
        "A browser-based tool that captures a photo, detects the ear, regresses landmarks, "
        "and outputs physical dimensions would replace the current workflow of either "
        "visiting an audiologist or guessing from a size chart. The full pipeline -- face "
        "detection for scale reference, ear detection, landmark regression, metric "
        "conversion -- already exists across the three projects in this series. Combining "
        "them into a single measurement tool is an engineering task, not a research one."
    )

    add_para(doc, "Model improvements", bold=True)
    add_para(
        doc,
        "Adaptive Wing loss (Wang et al., ICCV 2019) extends Wing loss with a "
        "per-sample power term that better handles varying difficulty across "
        "landmark points. It is implemented in the codebase but not yet evaluated. "
        "Attention mechanisms (channel or spatial) could improve fine localization "
        "at minimal parameter cost. Temporal smoothing for video sequences (Kalman "
        "filter or lightweight RNN over frame-to-frame predictions) would reduce "
        "jitter in the webcam demo."
    )

    add_para(doc, "Data expansion", bold=True)
    add_para(
        doc,
        "The current corpus is limited to 5,870 samples. Semi-supervised methods "
        "(pseudo-labeling ears detected in large face datasets like CelebA or FFHQ "
        "using the current model, then manually verifying a subset) could expand the "
        "training set by an order of magnitude. Cross-dataset evaluation on held-out "
        "sources would provide a stronger generalization assessment."
    )

    add_para(doc, "End-to-end pipeline", bold=True)
    add_para(
        doc,
        "Joint training of the detector and landmarker in a single differentiable "
        "pipeline could improve performance on difficult crops where the detector "
        "and landmarker have competing objectives. This is standard in face alignment "
        "(e.g., MTCNN, RetinaFace) but has not been explored for ears."
    )

    # ===== Technical Summary =====
    add_heading(doc, "Technical Summary")
    add_para(
        doc,
        "The codebase is 11 Python files and 2 JavaScript files. No external ML "
        "frameworks beyond PyTorch, Lightning, and ONNX Runtime. Training from "
        "preprocessed data to best checkpoint takes approximately 4 hours on a single "
        "consumer GPU. Inference runs at 579 FPS on GPU and in real time in the browser. "
        "The full browser demo requires no server -- two ONNX files, one HTML page, "
        "and one JavaScript module."
    )

    # ===== Prior Work =====
    add_heading(doc, "Prior Work in This Series")
    add_para_with_links(doc, [
        ("1. ", None),
        ("Trainable BlazeFace: High-Performance, Extendable Face Detection", BLAZEFACE_URL),
        (" -- Rebuilt MediaPipe's BlazeFace as a fully trainable PyTorch model. "
         "Established the BlazeBlock architecture, training pipeline, and ONNX browser "
         "deployment pattern used in all subsequent projects.", None),
    ])
    add_para_with_links(doc, [
        ("2. ", None),
        ("BlazeEar: Extending Trainable BlazeFace from Faces to Ears", BLAZEEAR_URL),
        (" -- Extended the BlazeFace architecture to detect ears. Same BlazeBlock "
         "backbone, anchor-based SSD detection head, trained from scratch on ear data. "
         "Provides the upstream detector for this project.", None),
    ])
    add_para_with_links(doc, [
        ("3. ", None),
        ("Measuring the Human Face Privately in the Browser", HEADSIZE_URL),
        (" -- Demonstrated metric head measurement from a single webcam image using "
         "facial landmarks and a reference object, all client-side. The ear sizing "
         "application described above is a direct extension of this approach.", None),
    ])

    # ===== Closing =====
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Built with PyTorch, Lightning, and a single RTX 5060 Ti. "
        "4 public datasets, 5,870 samples, 312K parameters. "
        "Third project in the BlazeFace/BlazeEar series."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


if __name__ == "__main__":
    doc = build_article()
    out = "C:/Users/shame/OneDrive/Desktop/ear_stuff/Ear_Landmarker/Ear_Landmarker_Article.docx"
    doc.save(out)
    print(f"Saved: {out}")
